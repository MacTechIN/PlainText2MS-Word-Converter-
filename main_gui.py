import re
import tkinter as tk
from tkinter import messagebox, scrolledtext
from docx import Document
from docx.shared import Pt, RGBColor
from docx.oxml.ns import qn
from docx.enum.style import WD_STYLE_TYPE


def save_as_docx(raw_text, filename):
    if not filename.endswith('.docx'):
        filename += '.docx'
    doc = Document()
    style = doc.styles['Normal']
    style.font.name = 'Malgun Gothic'
    style.element.rPr.rFonts.set(qn('w:eastAsia'), 'Malgun Gothic')
    style.font.size = Pt(10)

    # 제목 스타일 설정
    h1_style = doc.styles['Heading 1']
    h1_style.font.name = 'Malgun Gothic'
    h1_style.element.rPr.rFonts.set(qn('w:eastAsia'), 'Malgun Gothic')

    try:
        code_style = doc.styles.add_style('CodeStyle', WD_STYLE_TYPE.PARAGRAPH)
    except:
        code_style = doc.styles['CodeStyle']
    code_style.font.name = 'Consolas'
    code_style.element.rPr.rFonts.set(qn('w:eastAsia'), 'Malgun Gothic')
    code_style.font.size = Pt(10)

    lines = raw_text.split('\n')
    for line in lines:
        line = line.strip()
        if re.match(r'^.*Phase \d+:', line) or any(emoji in line for emoji in "🚀🛠️🗄️🔌🎨⚙️💻🔗📚✅"):
            doc.add_heading(line, level=1)
        elif re.match(r'^\d+\.\d+\s', line):
            doc.add_heading(line, level=2)
        elif not line:
            doc.add_paragraph("")
        else:
            p = doc.add_paragraph(line, style='CodeStyle')
            for run in p.runs:
                run.font.name = 'Consolas'
                run.element.rPr.rFonts.set(qn('w:eastAsia'), 'Malgun Gothic')
    doc.save(filename)
    return filename


def run_gui():
    window = tk.Tk()
    window.title("텍스트 to MS Word 변환기")
    window.geometry("600x700")

    # macOS에서 창이 바로 활성화되도록 강제 설정
    window.lift()
    window.attributes('-topmost', True)
    window.after_idle(window.attributes, '-topmost', False)
    window.focus_force()

    tk.Label(window, text="1. 내용을 아래에 붙여넣으세요:", font=("Malgun Gothic", 10, "bold")).pack(pady=5)
    text_area = scrolledtext.ScrolledText(window, wrap=tk.WORD, width=70, height=30)
    text_area.pack(padx=10, pady=5)
    text_area.focus_set()  # 입력창에 커서 바로 두기

    tk.Label(window, text="2. 저장할 파일명 입력:", font=("Malgun Gothic", 10, "bold")).pack(pady=5)
    filename_entry = tk.Entry(window, width=50)
    filename_entry.pack(pady=5)

    def start_conversion():
        input_text = text_area.get("1.0", tk.END).strip()
        input_filename = filename_entry.get().strip()
        if not input_text or not input_filename:
            messagebox.showwarning("경고", "내용과 파일명을 모두 입력해주세요.")
            return
        try:
            saved_name = save_as_docx(input_text, input_filename)
            messagebox.showinfo("성공", f"'{saved_name}' 저장 완료!")
        except Exception as e:
            messagebox.showerror("에러", f"오류 발생: {e}")

    tk.Button(window, text="워드 파일로 저장하기", command=start_conversion, bg="#4CAF50", fg="black").pack(pady=20)
    window.mainloop()


if __name__ == "__main__":
    run_gui()